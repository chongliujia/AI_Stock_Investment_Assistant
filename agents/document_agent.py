import os 
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from core.base_agent import BaseAgent
from core.llm_provider import OpenAIProvider
import re
from datetime import datetime

class DocumentAgent(BaseAgent):
    def __init__(self, output_dir="output_docs"):
        self.output_dir = output_dir
        self.llm_provider = OpenAIProvider()
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def _get_prompt_template(self, lang="en"):
        templates = {
            "en": {
                "create_doc": """Please create a {doc_type} document with approximately {word_count} words based on the following requirements:

                Topic: {topic}

                Please structure the document with:
                1. An engaging title
                2. An executive summary or introduction
                3. Main content with multiple sections and subheadings
                4. A conclusion
                
                Make sure the content is well-researched, informative, and engaging.
                Format the response with clear section breaks using '###' as separators."""
            },
            "zh": {
                "create_doc": """请创建一份详尽的{doc_type}，字数要求{word_count}字左右，基于以下要求：

                主题：{topic}

                请按照以下结构严格组织文档：
                1. 标题（简洁有力）
                2. 执行摘要（500字左右，概括文章主要观点）
                3. 引言（介绍背景、研究意义和主要问题）
                4. 主要内容（分4-5个主要章节展开）
                   - 每个章节需要2-3个子标题
                   - 包含具体数据支持
                   - 加入实际案例分析
                   - 引用权威研究或报告
                   - 讨论最新发展趋势
                5. 结论与展望（总结主要发现，并对未来发展做出预测）
                
                具体要求：
                1. 内容必须专业、深入且具有洞察力
                2. 每个观点都需要数据或案例支持
                3. 语言要专业流畅，避免空泛表述
                4. 适当加入图表描述（用文字描述图表内容）
                5. 确保逻辑性和连贯性
                
                使用'###'作为分节符来分隔不同部分。
                请确保生成足够详细的内容以达到字数要求。"""
            }
        }
        return templates.get(lang, templates["en"])["create_doc"]

    def _generate_filename(self, prompt, doc_type, lang="en"):
        # 从prompt中提取关键词作为文件名
        # 对于中文，取前10个字符；对于英文，取前5个单词
        if lang == "zh":
            # 提取中文字符
            chinese_chars = re.findall(r'[\u4e00-\u9fff]+', prompt)
            name_base = ''.join(chinese_chars)[:10]
        else:
            # 提取英文单词
            words = re.findall(r'\b\w+\b', prompt)
            name_base = '_'.join(words[:5]).lower()

        # 添加时间戳和文档类型
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        
        # 清理文件名中的非法字符
        name_base = re.sub(r'[\\/:*?"<>|]', '_', name_base)
        
        return f"{name_base}_{doc_type}_{timestamp}.docx"

    def _setup_document_style(self, doc, lang="en"):
        """设置文档的基本样式"""
        # 设置默认字体和大小
        style = doc.styles['Normal']
        if lang == "zh":
            style.font.name = 'SimSun'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        else:
            style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _add_title(self, doc, title, lang="en"):
        """添加文档标题"""
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        if lang == "zh":
            title_run.font.name = 'SimSun'
        else:
            title_run.font.name = 'Times New Roman'

    def handle_task(self, task):
        if task.task_type != "create_document":
            return f"Unsupported task type: {task.task_type}"

        # 获取任务参数
        doc_type = task.kwargs.get("doc_type", "general")
        word_count = task.kwargs.get("word_count", 1000)
        lang = task.kwargs.get("lang", "en")
        
        # 自动生成文件名（如果没有指定）
        filename = task.kwargs.get("filename")
        if not filename:
            filename = self._generate_filename(task.prompt, doc_type, lang)
        
        # 获取对应语言的提示词模板
        prompt_template = self._get_prompt_template(lang)
        prompt = prompt_template.format(
            doc_type=doc_type,
            word_count=word_count,
            topic=task.prompt
        )

        # 使用LLM生成内容
        content = self.llm_provider.generate_response(prompt)
        
        # 处理文件扩展名
        base_name, ext = os.path.splitext(filename)
        if not ext:
            filename = base_name + '.docx'
        
        file_path = os.path.join(self.output_dir, filename)

        try:
            # 创建Word文档
            doc = Document()
            
            # 设置文档基本样式
            self._setup_document_style(doc, lang)
            
            # 分割内容并格式化
            sections = content.split('###')
            first_section = True
            
            for section in sections:
                section = section.strip()
                if not section:
                    continue
                    
                lines = section.split('\n', 1)  # 只在第一个换行符处分割
                if not lines:
                    continue

                title = lines[0].strip()
                content = lines[1].strip() if len(lines) > 1 else ""

                if first_section:
                    # 第一个部分作为文档标题
                    self._add_title(doc, title, lang)
                    first_section = False
                else:
                    # 其他部分作为章节标题
                    heading = doc.add_heading(level=1)
                    heading_run = heading.add_run(title)
                    heading_run.font.size = Pt(14)
                    if lang == "zh":
                        heading_run.font.name = 'SimSun'

                if content:
                    # 添加段落内容
                    paragraph = doc.add_paragraph()
                    paragraph_run = paragraph.add_run(content)
                    if lang == "zh":
                        paragraph_run.font.name = 'SimSun'
                    
                    # 设置段落格式
                    paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
                    paragraph.paragraph_format.line_spacing = 1.5  # 1.5倍行距

            # 保存文档
            doc.save(file_path)
            return f"Document created successfully: {file_path}"
        except Exception as e:
            return f"Error creating document: {e}"

